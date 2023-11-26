from .models import *
from docx import Document
from .serializers import QuestionSerializer, AnswerSerializer
from django.core.exceptions import ObjectDoesNotExist
from rest_framework import permissions



def check_creator(telegram_id):
    if Creator.objects.filter(telegram_id=telegram_id):
        return True
    return False

def save_result(candidate_id, quiz_id, result, question_count):
    """Сохраняет результат прохождения тестирования пользователем в бд"""
    try:
        if int(result["result"]) > question_count:
            return "Error: number of questions should be greater or equal result"
        quiz = Quiz.objects.get(id=quiz_id)
        candidate = Candidate.objects.get(id=candidate_id)
        res = f'{int(result["result"])}/{question_count}'
        return Result.objects.create(quiz=quiz, candidate=candidate, result=res)
    except ObjectDoesNotExist as e:
        return f"Error: {str(e)}"


def filter_quiz(quiz):
    """Достает тест из бд и преобразует его в формат для json"""
    try:
        questions = Question.objects.filter(quiz=quiz)
        q_serializer = QuestionSerializer(questions, many=True)
    except ObjectDoesNotExist:
        return "Error: quiz was not found", None, None

    try:
        question_ids = questions.values_list('id', flat=True)
        answers = Answer.objects.filter(question__in=question_ids)
        a_serializer = AnswerSerializer(answers, many=True)
    except ObjectDoesNotExist:
        return "Error: questions was not found", None, None

    return q_serializer.data, a_serializer.data , len(question_ids)


def create_quiz_from_docx(file):
    """Парсит присланный docx файл и сохраняет его в бд"""
    try:
        doc = Document(file)
        if doc.paragraphs[0].style.name == 'Heading 1':
            quiz = Quiz.objects.create(title=doc.paragraphs[0].text)
        else:
            raise ValueError("Wrong file content")

        current_question = None
        for content in doc.paragraphs:
            if content.text == doc.paragraphs[0].text:
                continue
            if content.style.name == 'Heading 2':
                current_question = create_question(quiz, content.text)
            elif content.style.name == 'Heading 3':
                create_answer(current_question, content.text, True)
            else:
                create_answer(current_question, content.text, False)

        return "Quiz created successfully"

    except Exception as e:
        return f"Error: {str(e)}"


def create_question(quiz, content):
    if not quiz or not content:
        raise ValueError("Quiz and content are required to create a question")
    return Question.objects.create(quiz=quiz, content=content)


def create_answer(question, content, is_correct):
    if not question:
        raise ValueError("Missing question for answer")
    return Answer.objects.create(question=question, content=content, is_correct=is_correct)


class CreatorPermission(permissions.BasePermission):
    def has_permission(self, request, view):
        telegram_id = request.META.get('telegram_id')
        if not telegram_id:
            return False
        creator = Creator.objects.filter(telegram_id=telegram_id).exists()
        return creator
