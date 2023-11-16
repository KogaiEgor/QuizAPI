from docx import Document
from .models import Quiz, Answer, Question


def create_quiz_from_docx(file):
    try:
        doc = Document(file)
        if doc.paragraphs[0].style.name == 'Heading 1':
            quiz = Quiz.objects.create(title=doc.paragraphs[0].text)
        else:
            raise ValueError("Wrong file content")

        current_question = None
        for content in doc.paragraphs:
            if content.text == doc.paragraphs[0].text:
                continue
            if content.style.name == 'Heading 2':
                current_question = create_question(quiz, content.text)
            elif content.style.name == 'Heading 3':
                create_answer(current_question, content.text, True)
            else:
                create_answer(current_question, content.text, False)

        return "Quiz created successfully"

    except Exception as e:
        return f"Error: {str(e)}"


def create_question(quiz, content):
    return Question.objects.create(quiz=quiz, content=content)


def create_answer(question, content, is_correct):
    if not question:
        raise ValueError("Missing question for answer")
    return Answer.objects.create(question=question, content=content, is_correct=is_correct)

